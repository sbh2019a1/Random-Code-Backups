import xlrd
import random
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

loc = (r"C:\Users\beum3\PycharmProjects\excel2docx3\toefl_words.xlsx")

numofwords = 20
fontsize = 16

wb = xlrd.open_workbook(loc)
sheet = wb.sheet_by_index(0)
sheet.cell_value(0, 0)

wordList = []
ranomized30 = []

for i in range(sheet.nrows):
    wordList.append (sheet.cell_value(i, 0))

ranomizedWords = random.sample(wordList, numofwords)

print (ranomizedWords)
print (len(ranomizedWords))

#================export docx===================

document = Document()

a = document.add_paragraph('Word test : '+ str(numofwords) +' words')
b = document.add_paragraph('Name: ')

for run_a in a.runs:
    font = run_a.font
    font.size = Pt(20)
    font.name = 'Times New Roman'

table = document.add_table(rows=numofwords, cols=2)

i = 0
for i in range(numofwords):
    cell = table.cell(i, 0)
    cell.text = str(i+1) + '.  ' + ranomizedWords[i]

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(fontsize)
                font.name = 'Times New Roman'

table.style = 'TableGrid'
document.save('demo.docx')