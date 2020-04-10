import docx
from docx import Document
import random
from tkinter import *
from tkinter import scrolledtext

window = Tk()
window.title("Shuffling")
window.geometry('1000x800')

lbl1 = Label(window, text="Enter .docx file name (ex: shuffling_example.docx):  ")
lbl1.grid(column=0, row=0)

lbl2 = Label(window, text="Target file: ")
lbl2.grid(column=0, row=2)

txt = Entry(window,width=30)
txt.grid(column=0, row=1)

def clicked_filename_btn():
    textbox.delete(1.0,END)
    res = "Target file:" + txt.get()
    filename = txt.get()
    lbl2.configure(text= res)

    doc = docx.Document(filename)

    for x in range(len(doc.paragraphs)):
        text = doc.paragraphs[x].text
        text_split = text.split('.')
        if text_split[0]:
            text_split[0] = ' ' + text_split[0]

        random.shuffle(text_split)
        for y in text_split:
            if y:
                y = y + '.'
                textbox.insert(INSERT, y + '\n')
        textbox.insert(INSERT, ' '+'\n')

btn = Button(window, text="OK", command=clicked_filename_btn)
btn.grid(column=1, row=1)

textbox = scrolledtext.ScrolledText(window,width=120,height=50)
textbox.grid(column=0,row=3)

window.mainloop()
